"""Read text from all supported documents in a local folder, for use as a
fill source (alternative to RAG). Supports PDF, DOCX, XLSX, TXT, MD, HTML."""
import re
from pathlib import Path
import fitz  # PyMuPDF


def _read_pdf(p: Path) -> str:
    doc = fitz.open(str(p))
    try:
        if doc.needs_pass:
            return ""
        return "\n".join(page.get_text("text") for page in doc)
    finally:
        doc.close()


def _read_docx(p: Path) -> str:
    import docx
    d = docx.Document(str(p))
    parts = [par.text for par in d.paragraphs if par.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_xlsx(p: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    try:
        out = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    out.append("\t".join(cells))
        return "\n".join(out)
    finally:
        wb.close()


def _read_text(p: Path) -> str:
    return p.read_text(errors="ignore")


def _read_html(p: Path) -> str:
    return re.sub(r"<[^>]+>", " ", p.read_text(errors="ignore"))


READERS = {
    ".pdf": _read_pdf, ".docx": _read_docx, ".xlsx": _read_xlsx,
    ".txt": _read_text, ".md": _read_text, ".html": _read_html, ".htm": _read_html,
}


def read_folder_text(folder: Path, max_chars: int = 60000) -> tuple[str, list[str]]:
    """Return (combined_text, files_read). Skips hidden files and unreadable types."""
    parts: list[str] = []
    files_read: list[str] = []
    for f in sorted(folder.rglob("*")):
        if not f.is_file() or f.name.startswith(".") or f.suffix.lower() not in READERS:
            continue
        try:
            txt = READERS[f.suffix.lower()](f)
        except Exception:
            continue
        if txt and txt.strip():
            parts.append(f"### Source file: {f.name}\n{txt.strip()}")
            files_read.append(f.name)
    combined = "\n\n".join(parts)
    return combined[:max_chars], files_read
